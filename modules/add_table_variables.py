import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_table_from_dataframe(
    doc: Document, df: pd.DataFrame, font_size: int = 9
):
    for paragraph in doc.paragraphs:
        if 'var_table' in paragraph.text:
            # Удалить текст 'var_table' из найденного параграфа:
            paragraph.text = paragraph.text.replace('var_table', '')

            # Создать таблицу в месте, где найден 'var_table':
            table = doc.add_table(rows=1, cols=len(df.columns))

            # Настройка стиля таблицы:
            tbl = table._element
            tblBorders = parse_xml(
                r'<w:tblBorders %s>'
                r'<w:top w:val="single" w:sz="4" w:space="0"/>'
                r'<w:left w:val="single" w:sz="4" w:space="0"/>'
                r'<w:bottom w:val="single" w:sz="4" w:space="0"/>'
                r'<w:right w:val="single" w:sz="4" w:space="0"/>'
                r'<w:insideH w:val="single" w:sz="4" w:space="0"/>'
                r'<w:insideV w:val="single" w:sz="4" w:space="0"/>'
                r'</w:tblBorders>' % nsdecls('w')
            )
            tblPr = tbl.find(
                './/w:tblPr',
                namespaces={
                    'w':
                    'http://schemas.openxmlformats.org' +
                    '/wordprocessingml/2006/main'
                }
            )
            if tblPr is None:
                tblPr = parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
                tbl.append(tblPr)
            tblPr.append(tblBorders)

            # Центрирование таблицы в документе:
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Добавление заголовков:
            header_cells = table.rows[0].cells
            for i, column_name in enumerate(df.columns):
                cell = header_cells[i]
                cell.text = column_name

            # Форматирование заголовков:
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(font_size)
                    run.font.name = 'Times New Roman'

            # Добавление данных:
            for i, row in df.iterrows():
                row_cells = table.add_row().cells
                for j, value in enumerate(row):
                    cell = row_cells[j]
                    cell.text = str(value)

            # Форматирование всех ячеек данных:
            for row in table.rows[1:]:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.runs[0]
                        run.font.size = Pt(font_size)
                        run.font.name = 'Times New Roman'

            break
