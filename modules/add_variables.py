import os
import shutil

import pandas as pd
from typing import Optional
from docxtpl import DocxTemplate
from docx import Document

from modules.add_table_variables import add_table_from_dataframe


def add_variables(
    context: dict,
    template_path: str,
    output_name: str,
    output_dir: str,
    table: Optional[pd.DataFrame] = None
):
    """
    Шаблон сначала рендерится с переменными.
    Создается временный файл с отрендеренным шаблоном.
    Затем временный файл загружается с помощью python-docx, и в него
    добавляется таблица. Финальный документ сохраняется, а временный файл
    удаляется.
    """
    output_path = os.path.join(output_dir, output_name)

    # Создание копии шаблона:
    shutil.copy(template_path, output_path)

    # Загрузка копии шаблона Word:
    doc: Document = DocxTemplate(output_path)

    # Рендеринг документа с заменой переменных:
    doc.render(context)

    # Сохранение временного документа:
    temp_path = os.path.join(output_dir, 'temp.docx')
    doc.save(temp_path)

    # Загрузка временного документа для добавления таблицы:
    final_doc = Document(temp_path)

    # Добавляем таблицу в документ:
    if table is not None:
        add_table_from_dataframe(final_doc, table)

    # Сохранение измененного документа:
    final_doc.save(output_path)

    # Удаление временного документа:
    os.remove(temp_path)
